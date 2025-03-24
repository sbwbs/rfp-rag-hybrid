from typing import List, Dict, BinaryIO
import os
import logging
import tempfile
import PyPDF2
from docx import Document as DocxDocument
import pandas as pd
from pathlib import Path

# Set up logging
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles document text extraction from various file formats."""
    
    def extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF files."""
        try:
            logger.info(f"Extracting text from PDF: {file_path}")
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            
            logger.info(f"Successfully extracted {len(text)} characters from PDF")
            return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            raise
    
    def extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX files."""
        try:
            logger.info(f"Extracting text from DOCX: {file_path}")
            doc = DocxDocument(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            logger.info(f"Successfully extracted {len(text)} characters from DOCX")
            return text
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise
    
    def extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT files."""
        try:
            logger.info(f"Extracting text from TXT: {file_path}")
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            logger.info(f"Successfully extracted {len(text)} characters from TXT")
            return text
        except Exception as e:
            logger.error(f"Error extracting text from TXT: {e}")
            raise
    
    def extract_from_xlsx(self, file_path: str) -> str:
        """Extract text from XLSX files."""
        try:
            logger.info(f"Extracting text from XLSX: {file_path}")
            # Read all sheets
            xlsx = pd.ExcelFile(file_path)
            text = ""
            
            for sheet_name in xlsx.sheet_names:
                df = pd.read_excel(xlsx, sheet_name=sheet_name)
                text += f"Sheet: {sheet_name}\n"
                text += df.to_string(index=False) + "\n\n"
            
            logger.info(f"Successfully extracted {len(text)} characters from XLSX")
            return text
        except Exception as e:
            logger.error(f"Error extracting text from XLSX: {e}")
            raise
    
    def extract_text(self, file_path: str) -> str:
        """Main method to extract text from any supported file type."""
        try:
            logger.info(f"Processing file: {file_path}")
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension == '.pdf':
                return self.extract_from_pdf(file_path)
            elif file_extension == '.docx':
                return self.extract_from_docx(file_path)
            elif file_extension == '.txt':
                return self.extract_from_txt(file_path)
            elif file_extension in ['.xlsx', '.xls']:
                return self.extract_from_xlsx(file_path)
            else:
                error_msg = f"Unsupported file format: {file_extension}"
                logger.error(error_msg)
                raise ValueError(error_msg)
        except Exception as e:
            logger.error(f"Error in extract_text: {e}")
            raise
    
    def process_uploaded_file(self, uploaded_file: BinaryIO) -> str:
        """Process an uploaded file from Streamlit."""
        try:
            logger.info(f"Processing uploaded file: {uploaded_file.name}")
            # Create a temporary file
            temp_dir = tempfile.mkdtemp()
            temp_path = os.path.join(temp_dir, uploaded_file.name)
            
            # Write the uploaded file to a temporary location
            with open(temp_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            # Extract text
            text = self.extract_text(temp_path)
            
            # Clean up the temporary file
            try:
                os.remove(temp_path)
                os.rmdir(temp_dir)
            except:
                logger.warning(f"Failed to clean up temporary file: {temp_path}")
            
            return text
        except Exception as e:
            logger.error(f"Error processing uploaded file: {e}")
            raise
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Split text into overlapping chunks."""
        try:
            logger.info(f"Chunking text of length {len(text)} with chunk size {chunk_size} and overlap {overlap}")
            chunks = []
            start = 0
            
            while start < len(text):
                end = min(start + chunk_size, len(text))
                # If we're not at the end of the text, try to find a good break point
                if end < len(text):
                    # Try to break at paragraph
                    paragraph_break = text.rfind('\n\n', start, end)
                    if paragraph_break != -1 and paragraph_break > start + 0.5 * chunk_size:
                        end = paragraph_break
                    else:
                        # Try to break at sentence
                        sentence_break = text.rfind('. ', start, end)
                        if sentence_break != -1 and sentence_break > start + 0.5 * chunk_size:
                            end = sentence_break + 1  # Include the period
                
                chunks.append(text[start:end])
                start = end - overlap
            
            logger.info(f"Created {len(chunks)} chunks")
            return chunks
        except Exception as e:
            logger.error(f"Error chunking text: {e}")
            raise